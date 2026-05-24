from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)

def add_centered(text, bold=False, size=12, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_justified(text, bold=False, size=12, indent=False, space_before=0, space_after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_blank(space=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)

# CABECALHO
add_centered("EXCELENTISSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA", size=12)
add_centered("147 VARA CIVEL DO FORO CENTRAL DA COMARCA DE CURITIBA", bold=True, size=12)
add_centered("ESTADO DO PARANA", size=12, space_after=18)

add_centered("Autos n. 0000295-45.2025.8.16.0194", bold=True, size=12, space_after=4)
add_centered("Autora: AZUL COMPANHIA DE SEGUROS GERAIS", size=12)
add_centered("Reu: SELMO [SOBRENOME A CONFIRMAR]", size=12, space_after=24)

# TITULO
add_centered("MINUTA DE ACORDO PARA QUITACAO DE", bold=True, size=13, space_after=2)
add_centered("HONORARIOS ADVOCATICIOS SUCUMBENCIAIS", bold=True, size=13, space_after=20)

# PREAMBULO
add_justified(
    "SELMO [SOBRENOME A CONFIRMAR], ja qualificado nos autos, REU nesta acao, "
    "por seu advogado infra-assinado, vem respeitosamente a presenca de Vossa "
    "Excelencia, em face de AZUL COMPANHIA DE SEGUROS GERAIS, AUTORA nesta acao, "
    "apresentar a presente MINUTA DE ACORDO para formalizacao da proposta de "
    "transacao extrajudicial oferecida pela propria AUTORA, referente aos honorarios "
    "advocaticios sucumbenciais fixados nos presentes autos, nos termos dos "
    "arts. 840 a 850 do Codigo Civil e do art. 515, III, do Codigo de Processo Civil, "
    "mediante as clausulas e condicoes a seguir:",
    indent=True, space_after=16
)

# CL. 1
add_justified("CLAUSULA PRIMEIRA - DO OBJETO", bold=True, space_before=4, space_after=4)
add_justified(
    "O presente instrumento tem por objeto a transacao dos honorarios advocaticios "
    "sucumbenciais devidos ao advogado do REU, fixados pela r. sentenca de "
    "17/05/2026, que julgou IMPROCEDENTE a acao, condenando a AUTORA ao pagamento "
    "de honorarios sucumbenciais correspondentes a 10% (dez por cento) sobre o "
    "valor atualizado da causa, nos termos do art. 85, par. 2, do CPC.",
    indent=True, space_after=12
)

# CL. 2
add_justified("CLAUSULA SEGUNDA - DO VALOR E DA QUITACAO", bold=True, space_before=4, space_after=4)
add_justified(
    "As partes transacionam os honorarios sucumbenciais pelo valor de "
    "R$ 6.000,00 (seis mil reais), montante que, uma vez pago integralmente "
    "e no prazo estipulado neste instrumento, importara plena, geral, irrevogavel "
    "e irretratavel quitacao de toda e qualquer obrigacao da AUTORA a titulo de "
    "honorarios advocaticios sucumbenciais decorrentes dos presentes autos.",
    indent=True, space_after=12
)

# CL. 3
add_justified("CLAUSULA TERCEIRA - DO PRAZO E DA FORMA DE PAGAMENTO", bold=True, space_before=4, space_after=4)
add_justified(
    "O pagamento do valor de R$ 6.000,00 (seis mil reais) devera ser efetuado "
    "pela AUTORA em parcela unica, no prazo improrrogavel de 10 (dez) dias "
    "corridos, contados da data de protocolo da presente minuta nos autos, "
    "mediante transferencia via PIX para os dados bancarios indicados a seguir:",
    indent=True, space_after=8
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(8)
p.paragraph_format.left_indent  = Cm(2)
p.paragraph_format.right_indent = Cm(2)
run = p.add_run(
    "Beneficiario: JUST HELP\n"
    "CNPJ: 54.790.349/0001-01\n"
    "Banco: 336 - Banco C6 S.A.\n"
    "Agencia: 0001  |  Conta Corrente: 32709981-0\n"
    "Chave PIX (e-mail): c6@justhelp.com.br"
)
run.font.size = Pt(11)
run.bold = True

add_blank(space=6)
add_justified(
    "Paragrafo unico. O comprovante de pagamento devera ser encaminhado ao "
    "advogado do REU no prazo de ate 24 (vinte e quatro) horas apos a realizacao "
    "da transferencia, sob pena de nao reconhecimento da quitacao.",
    indent=True, space_after=12
)

# CL. 4 - CLAUSULA PENAL
add_justified("CLAUSULA QUARTA - DA CLAUSULA PENAL E DO INADIMPLEMENTO", bold=True, space_before=4, space_after=4)
add_justified(
    "Em caso de inadimplemento da obrigacao de pagamento no prazo e na forma "
    "estabelecidos na Clausula Terceira, o presente acordo perdera "
    "automaticamente seus efeitos, independentemente de notificacao ou "
    "interpelacao judicial ou extrajudicial, ficando restabelecidos "
    "integralmente todos os direitos do advogado do REU de promover o "
    "cumprimento de sentenca pelo valor integral dos honorarios sucumbenciais "
    "fixados em sentenca, devidamente atualizado, acrescido de:",
    indent=True, space_after=8
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.left_indent = Cm(1.25)
run = p.add_run(
    "a) multa moratoria de 10% (dez por cento) sobre o valor total dos honorarios apurados na data do inadimplemento;\n"
    "b) juros de mora de 1% (um por cento) ao mes, pro rata die, desde a data do vencimento;\n"
    "c) correcao monetaria pelo IPCA-E, desde o transito em julgado da sentenca;\n"
    "d) honorarios advocaticios adicionais de 10% (dez por cento) sobre o total executado, referentes a fase de cumprimento de sentenca."
)
run.font.size = Pt(12)

add_blank(space=4)
add_justified(
    "Paragrafo unico. A aceitacao do pagamento em atraso nao implica novacao "
    "ou renuncia as penalidades previstas nesta clausula, salvo declaracao "
    "expressa em contrario pelo advogado do REU.",
    indent=True, space_after=12
)

# CL. 5
add_justified("CLAUSULA QUINTA - DOS EFEITOS DO PAGAMENTO", bold=True, space_before=4, space_after=4)
add_justified(
    "Efetuado o pagamento integral e confirmada a transferencia, o advogado do "
    "REU compromete-se a peticionar nos autos, no prazo de 5 (cinco) dias uteis, "
    "requerendo a extincao do feito com resolucao do merito, nos termos do "
    "art. 487, III, alínea b, do CPC, exclusivamente quanto aos honorarios "
    "sucumbenciais objeto deste acordo.",
    indent=True, space_after=12
)

# CL. 6
add_justified("CLAUSULA SEXTA - DA HOMOLOGACAO JUDICIAL", bold=True, space_before=4, space_after=4)
add_justified(
    "As partes requerem a Vossa Excelencia a homologacao do presente acordo, "
    "nos termos do art. 515, III, c/c art. 725, VIII, do CPC, com a "
    "consequente extincao do processo com resolucao do merito, apos comprovacao "
    "do pagamento nos autos.",
    indent=True, space_after=12
)

# CL. 7
add_justified("CLAUSULA SETIMA - DO FORO", bold=True, space_before=4, space_after=4)
add_justified(
    "As partes elegem o foro da Comarca de Curitiba/PR para dirimir quaisquer "
    "controversias decorrentes do presente instrumento, com renuncia expressa "
    "a qualquer outro, por mais privilegiado que seja.",
    indent=True, space_after=18
)

add_justified("Nesses termos, requerem a homologacao.", space_before=4, space_after=20)

add_centered("Curitiba, _____ de __________________ de 2026.", space_after=40)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run(
    "_________________________________________          _________________________________________\n"
    "Advogado(a) da AUTORA                                           Advogado(a) do REU - GCJ ADVOCACIA\n"
    "AZUL COMPANHIA DE SEGUROS GERAIS                  SELMO [SOBRENOME]\n"
    "OAB/PR n. __________                                               OAB/PR n. __________"
)
run.font.size = Pt(11)

add_blank(space=18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Processo n. 0000295-45.2025.8.16.0194  |  147 Vara Civel de Curitiba/PR\n"
    "AZUL COMPANHIA DE SEGUROS GERAIS  x  SELMO [SOBRENOME]"
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save("y:/Projeto Claude/Minuta_Acordo_Honorarios_Selmo.docx")
print("Documento salvo com sucesso.")
